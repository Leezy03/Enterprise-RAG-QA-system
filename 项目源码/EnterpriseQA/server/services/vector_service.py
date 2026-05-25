"""
文档向量化服务
负责文档解析、文本分块和Chroma向量存储
"""
import os
import time
import logging

os.environ.setdefault('ANONYMIZED_TELEMETRY', 'False')
logging.getLogger('chromadb.telemetry.product.posthog').disabled = True

from flask import current_app
from ollama import Client as OllamaClient, ResponseError
from chromadb.config import Settings
from langchain_core.documents import Document as LangChainDocument
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from services.bm25_service import BM25Service
from services.rerank_service import RerankService


class OllamaServiceError(Exception):
    """Ollama服务相关异常，用于提供更清晰的错误提示"""
    pass


class VectorService:
    """文档向量化服务类"""

    def __init__(self):
        """初始化嵌入模型和文本分割器"""
        self.embeddings = OllamaEmbeddings(
            model=current_app.config['OLLAMA_EMBED_MODEL'],
            base_url=current_app.config['OLLAMA_BASE_URL']
        )
        self.default_chunk_strategy = {
            'chunk_size': current_app.config['CHUNK_SIZE'],
            'chunk_overlap': current_app.config['CHUNK_OVERLAP']
        }
        self.chunk_strategies = current_app.config.get('CHUNK_STRATEGIES', {})
        splitter_file_types = {'txt', 'pdf', 'md', 'docx'} | set(self.chunk_strategies.keys())
        self.text_splitters = {
            file_type: self._create_text_splitter(self._resolve_chunk_strategy(file_type))
            for file_type in splitter_file_types
        }
        self.text_splitter = self._create_text_splitter(self.default_chunk_strategy)
        self.markdown_headers_to_split_on = [
            ('#', 'header_1'),
            ('##', 'header_2'),
            ('###', 'header_3'),
            ('####', 'header_4'),
        ]
        self.markdown_header_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=self.markdown_headers_to_split_on,
            strip_headers=False
        )
        self.persist_dir = current_app.config['CHROMA_PERSIST_DIR']
        self.batch_size = current_app.config.get('EMBED_BATCH_SIZE', 10)
        self.max_retries = current_app.config.get('EMBED_MAX_RETRIES', 3)
        self.bm25_service = BM25Service(
            k1=current_app.config.get('BM25_K1', 1.5),
            b=current_app.config.get('BM25_B', 0.75)
        )

    def _resolve_chunk_strategy(self, file_type):
        """获取指定文件类型的分块策略，并兜底处理异常配置。"""
        strategy = dict(self.default_chunk_strategy)
        strategy.update(self.chunk_strategies.get(file_type, {}))

        chunk_size = int(strategy.get('chunk_size') or self.default_chunk_strategy['chunk_size'])
        chunk_overlap = int(strategy.get('chunk_overlap') or 0)
        if chunk_overlap >= chunk_size:
            current_app.logger.warning(
                f'{file_type} chunk_overlap({chunk_overlap}) 不能大于等于 chunk_size({chunk_size})，已自动调整'
            )
            chunk_overlap = max(0, chunk_size // 10)

        strategy['chunk_size'] = chunk_size
        strategy['chunk_overlap'] = chunk_overlap
        return strategy

    def _create_text_splitter(self, strategy):
        """根据分块策略创建递归字符切分器。"""
        splitter_kwargs = {
            'chunk_size': strategy['chunk_size'],
            'chunk_overlap': strategy['chunk_overlap'],
            'length_function': len
        }
        separators = strategy.get('separators')
        if separators:
            splitter_kwargs['separators'] = separators
        return RecursiveCharacterTextSplitter(**splitter_kwargs)

    def _get_text_splitter(self, file_type):
        """按文件类型获取切分器，未知类型使用默认切分器。"""
        return self.text_splitters.get(file_type, self.text_splitter)

    def _get_chunk_strategy_metadata(self, file_type):
        """把实际切分参数写入metadata，便于调试和来源追踪。"""
        strategy = self._resolve_chunk_strategy(file_type)
        return {
            'chunk_strategy': file_type if file_type in self.chunk_strategies else 'default',
            'chunk_size': strategy['chunk_size'],
            'chunk_overlap': strategy['chunk_overlap']
        }

    def _get_offset_match_text(self, chunk):
        """去掉人为补充的上下文前缀，用原文片段定位字符范围。"""
        if chunk.startswith('标题路径：') and '\n' in chunk:
            return chunk.split('\n', 1)[1].strip()
        return chunk.strip()

    def _find_chunk_offsets(self, source_text, chunks):
        """
        在源文档文本中定位每个chunk的字符范围。
        该范围用于source tracking，不参与检索；找不到时跳过，避免影响入库。
        """
        offsets = []
        cursor = 0
        for chunk in chunks:
            match_text = self._get_offset_match_text(chunk)
            if not match_text:
                offsets.append({})
                continue

            start = source_text.find(match_text, cursor)
            if start == -1:
                start = source_text.find(match_text)
            if start == -1:
                offsets.append({})
                continue

            end = start + len(match_text)
            offsets.append({
                'char_start': start,
                'char_end': end
            })
            cursor = end

        return offsets

    def _check_ollama(self):
        """
        预检查Ollama服务可用性和嵌入模型是否就绪。
        仅在"服务未启动"和"模型未安装"时硬拦截；
        5xx等瞬时错误只记录警告，让后续重试机制处理。
        :raises OllamaServiceError: 服务不可达或模型未安装时抛出
        """
        base_url = current_app.config['OLLAMA_BASE_URL']
        model_name = current_app.config['OLLAMA_EMBED_MODEL']
        client = OllamaClient(host=base_url)

        try:
            model_list = client.list()
        except ConnectionError:
            raise OllamaServiceError(
                f'无法连接Ollama服务({base_url})，请确认Ollama已启动'
            )
        except ResponseError as e:
            current_app.logger.warning(
                f'Ollama预检查返回异常(status {e.status_code})，将继续尝试向量化: {e}'
            )
            return
        except Exception as e:
            current_app.logger.warning(f'Ollama预检查失败，将继续尝试向量化: {e}')
            return

        installed = {m.get('name', '') for m in model_list.get('models', [])}
        if not any(model_name in name or name in model_name for name in installed):
            raise OllamaServiceError(
                f'嵌入模型 {model_name} 未安装，请先执行: ollama pull {model_name}'
            )

    def _get_collection_name(self, kb_id):
        """
        根据知识库ID生成Chroma集合名称
        每个知识库使用独立的collection进行隔离
        """
        return f"kb_{kb_id}"

    def _get_vectorstore(self, kb_id):
        """
        获取指定知识库对应的Chroma向量库实例
        :param kb_id: 知识库ID
        :return: Chroma向量库
        """
        return Chroma(
            collection_name=self._get_collection_name(kb_id),
            embedding_function=self.embeddings,
            persist_directory=self.persist_dir,
            client_settings=Settings(anonymized_telemetry=False)
        )

    def _load_documents(self, file_path, file_type):
        """
        根据文件类型加载文档内容，并保留可用于来源追踪的结构化metadata。
        :param file_path: 文件路径
        :param file_type: 文件类型（txt/pdf/md/docx）
        :return: LangChain Document列表
        """
        if file_type in ('txt', 'md'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return [
                LangChainDocument(
                    page_content=text,
                    metadata={'file_type': file_type}
                )
            ]

        if file_type == 'pdf':
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            documents = []
            for page_index, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    documents.append(
                        LangChainDocument(
                            page_content=page_text,
                            metadata={
                                'file_type': file_type,
                                'page_number': page_index
                            }
                        )
                    )
            return documents

        if file_type == 'docx':
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            documents = []
            section_lines = []
            section_title = None
            section_index = 0
            paragraph_start = None
            paragraph_end = None

            def flush_section():
                nonlocal section_lines, section_title, section_index, paragraph_start, paragraph_end
                if not section_lines:
                    return
                metadata = {
                    'file_type': file_type,
                    'section_index': section_index,
                    'paragraph_start': paragraph_start,
                    'paragraph_end': paragraph_end
                }
                if section_title:
                    metadata['section_title'] = section_title
                documents.append(
                    LangChainDocument(
                        page_content='\n'.join(section_lines),
                        metadata=metadata
                    )
                )
                section_lines = []
                paragraph_start = None
                paragraph_end = None

            for paragraph_index, para in enumerate(doc.paragraphs, 1):
                text = para.text.strip()
                if not text:
                    continue

                style_name = getattr(para.style, 'name', '') or ''
                is_heading = style_name.lower().startswith('heading')
                if is_heading and section_lines:
                    flush_section()

                if is_heading:
                    section_index += 1
                    section_title = text
                elif section_index == 0:
                    section_index = 1

                if paragraph_start is None:
                    paragraph_start = paragraph_index
                paragraph_end = paragraph_index
                section_lines.append(text)

            flush_section()
            return documents

        return []

    def _split_text(self, text, file_type):
        """
        根据文件类型切分文本。
        Markdown优先按标题层级切分，再对过长标题块做递归字符切分。
        """
        if file_type == 'md':
            return self._split_markdown_text(text)

        splitter = self._get_text_splitter(file_type)
        chunks = splitter.split_text(text)
        return chunks, [self._get_chunk_strategy_metadata(file_type) for _ in chunks]

    def _split_markdown_text(self, text):
        """
        使用Markdown标题结构切分文本，并把标题路径写入metadata。
        标题路径也会加入chunk文本，提升针对章节标题类问题的召回稳定性。
        """
        header_docs = self.markdown_header_splitter.split_text(text)
        if not header_docs:
            splitter = self._get_text_splitter('md')
            chunks = splitter.split_text(text)
            return chunks, [self._get_chunk_strategy_metadata('md') for _ in chunks]

        chunks = []
        chunk_metadatas = []
        header_keys = [name for _, name in self.markdown_headers_to_split_on]
        splitter = self._get_text_splitter('md')

        for header_doc in header_docs:
            header_metadata = {
                key: value
                for key, value in (header_doc.metadata or {}).items()
                if value
            }
            header_path = ' / '.join(
                str(header_metadata[key])
                for key in header_keys
                if header_metadata.get(key)
            )
            if header_path:
                header_metadata['header_path'] = header_path

            sub_chunks = splitter.split_text(header_doc.page_content)
            for sub_chunk in sub_chunks:
                chunk_text = sub_chunk
                if header_path and header_path not in sub_chunk[:120]:
                    chunk_text = f"标题路径：{header_path}\n{sub_chunk}"
                chunks.append(chunk_text)
                chunk_metadata = self._get_chunk_strategy_metadata('md')
                chunk_metadata.update(header_metadata)
                chunk_metadatas.append(chunk_metadata)

        return chunks, chunk_metadatas

    def _add_texts_with_retry(self, vectorstore, texts, metadatas, ids):
        """
        带重试的向量写入，处理Ollama瞬时故障(502/503/504等)
        :param vectorstore: Chroma向量库实例
        :param texts: 文本分块列表
        :param metadatas: 元数据列表
        :param ids: ID列表
        """
        last_error = None
        for attempt in range(self.max_retries):
            try:
                vectorstore.add_texts(texts=texts, metadatas=metadatas, ids=ids)
                return
            except Exception as e:
                last_error = e
                err_msg = str(e)
                is_retryable = any(code in err_msg for code in ('502', '503', '504'))
                if not is_retryable or attempt == self.max_retries - 1:
                    raise
                wait = 2 ** attempt
                current_app.logger.warning(
                    f'Ollama嵌入请求失败(第{attempt + 1}次)，{wait}秒后重试: {err_msg}'
                )
                time.sleep(wait)
        raise last_error

    def process_document(self, doc_id, file_path, file_type, kb_id, original_file_name=None):
        """
        处理文档：预检查 -> 解析文件 -> 文本分块 -> 分批存入向量库
        :param doc_id: 文档ID
        :param file_path: 文件路径
        :param file_type: 文件类型
        :param kb_id: 知识库ID
        :param original_file_name: 用户上传时的原始文件名
        :return: 分块数量
        """
        self._check_ollama()

        source_documents = self._load_documents(file_path, file_type)
        if not any(doc.page_content.strip() for doc in source_documents):
            raise ValueError('文档内容为空，无法进行向量化')

        chunks = []
        chunk_metadatas = []
        for source_doc_index, source_doc in enumerate(source_documents, 1):
            source_text = source_doc.page_content
            if not source_text.strip():
                continue
            source_chunks, source_chunk_metadatas = self._split_text(source_text, file_type)
            source_chunk_offsets = self._find_chunk_offsets(source_text, source_chunks)
            for i, chunk in enumerate(source_chunks):
                source_metadata = dict(source_doc.metadata or {})
                split_metadata = source_chunk_metadatas[i] if i < len(source_chunk_metadatas) else {}
                offset_metadata = source_chunk_offsets[i] if i < len(source_chunk_offsets) else {}
                chunks.append(chunk)
                chunk_metadatas.append({
                    'source_doc_index': source_doc_index,
                    **source_metadata,
                    **split_metadata,
                    **offset_metadata
                })

        if not chunks:
            raise ValueError('文档分块失败')

        file_name = original_file_name or os.path.basename(file_path)
        metadatas = [
            {
                'doc_id': doc_id,
                'file_name': file_name,
                'stored_file_name': os.path.basename(file_path),
                'file_type': file_type,
                'chunk_index': i,
                **chunk_metadatas[i]
            }
            for i in range(len(chunks))
        ]
        ids = [f"doc_{doc_id}_chunk_{i}" for i in range(len(chunks))]

        vectorstore = self._get_vectorstore(kb_id)

        # 分批写入，降低单次Ollama嵌入请求的压力
        for i in range(0, len(chunks), self.batch_size):
            batch_end = min(i + self.batch_size, len(chunks))
            self._add_texts_with_retry(
                vectorstore,
                texts=chunks[i:batch_end],
                metadatas=metadatas[i:batch_end],
                ids=ids[i:batch_end],
            )

        return len(chunks)

    def delete_document(self, doc_id, kb_id):
        """
        从向量库中删除指定文档的所有分块
        :param doc_id: 文档ID
        :param kb_id: 知识库ID
        """
        vectorstore = self._get_vectorstore(kb_id)
        # 根据文档ID过滤并删除
        vectorstore._collection.delete(where={'doc_id': doc_id})

    def _candidate_key(self, doc):
        """生成跨检索方式去重用的稳定键。"""
        metadata = doc.metadata or {}
        doc_id = metadata.get('doc_id')
        chunk_index = metadata.get('chunk_index')
        if doc_id is not None and chunk_index is not None:
            return ('doc_chunk', str(doc_id), str(chunk_index))

        chroma_id = metadata.get('chroma_id')
        if chroma_id:
            return ('chroma_id', str(chroma_id))

        return ('content', doc.page_content or '')

    def _vector_search(self, vectorstore, query, initial_k):
        """执行向量召回，并保留向量相似度与召回排名。"""
        try:
            docs_with_scores = vectorstore.similarity_search_with_score(query, k=initial_k)
            docs = []
            for rank, (doc, distance) in enumerate(docs_with_scores, 1):
                doc.metadata = dict(doc.metadata or {})
                distance = max(float(distance), 0)
                normalized_score = 1 / (1 + distance)
                doc.metadata['vector_score'] = round(normalized_score, 4)
                doc.metadata['relevance_score'] = round(normalized_score, 4)
                doc.metadata['vector_rank'] = rank
                doc.metadata['retrieval_method'] = 'vector'
                docs.append(doc)
            return docs
        except Exception as e:
            current_app.logger.warning(f'带分数检索失败，降级为普通向量检索: {e}')
            docs = vectorstore.similarity_search(query, k=initial_k)
            for rank, doc in enumerate(docs, 1):
                doc.metadata = dict(doc.metadata or {})
                doc.metadata['vector_score'] = 0
                doc.metadata['relevance_score'] = 0
                doc.metadata['vector_rank'] = rank
                doc.metadata['retrieval_method'] = 'vector'
            return docs

    def _build_bm25_search_text(self, content, metadata):
        """拼接BM25可检索字段，保留正文为主，metadata用于文件名/标题类问题召回。"""
        if not current_app.config.get('BM25_INCLUDE_METADATA', True):
            return content

        metadata = metadata or {}
        fields = [content or '']
        for field in ('file_name', 'stored_file_name', 'header_path', 'section_title', 'file_type'):
            value = metadata.get(field)
            if value:
                fields.append(str(value))
        return '\n'.join(fields)

    def _bm25_search(self, vectorstore, query, top_k):
        """
        在当前Chroma集合内执行BM25词法召回。
        BM25用于补足向量检索对专有名词、编号、制度条款、英文缩写的精确匹配不足。
        """
        scan_limit = current_app.config.get('HYBRID_KEYWORD_SCAN_LIMIT', 2000)
        get_kwargs = {'include': ['documents', 'metadatas']}
        if scan_limit and scan_limit > 0:
            get_kwargs['limit'] = scan_limit

        try:
            raw = vectorstore._collection.get(**get_kwargs)
        except Exception as e:
            current_app.logger.warning(f'BM25召回失败，将仅使用向量召回: {e}')
            return []

        documents = raw.get('documents') or []
        metadatas = raw.get('metadatas') or []
        ids = raw.get('ids') or []
        corpus = []

        for index, content in enumerate(documents):
            if not content:
                continue

            metadata = dict(metadatas[index] or {}) if index < len(metadatas) else {}
            if index < len(ids):
                metadata['chroma_id'] = ids[index]

            corpus.append({
                'content': content,
                'search_text': self._build_bm25_search_text(content, metadata),
                'metadata': metadata
            })

        ranked = self.bm25_service.rank(query, corpus, top_k)
        candidates = []
        for rank, item in enumerate(ranked, 1):
            metadata = dict(item['metadata'])
            bm25_score = round(float(item['bm25_score']), 4)
            bm25_score_raw = round(float(item['bm25_score_raw']), 4)

            metadata['bm25_score'] = bm25_score
            metadata['bm25_score_raw'] = bm25_score_raw
            metadata['bm25_query_terms'] = item['bm25_query_terms']
            metadata['bm25_matched_terms'] = item['bm25_matched_terms']
            metadata['bm25_matched_keywords'] = item['bm25_matched_keywords']
            metadata['keyword_score'] = bm25_score
            metadata['keyword_search_score'] = bm25_score
            metadata['keyword_backend'] = 'bm25'
            metadata['relevance_score'] = bm25_score
            metadata['keyword_rank'] = rank
            metadata['retrieval_method'] = 'bm25'
            candidates.append(
                LangChainDocument(
                    page_content=item['content'],
                    metadata=metadata
                )
            )

        return candidates

    def _keyword_search(self, vectorstore, query, top_k):
        """Backward-compatible wrapper. The keyword branch is now BM25."""
        return self._bm25_search(vectorstore, query, top_k)

    def _merge_hybrid_results(self, vector_docs, keyword_docs):
        """使用RRF融合向量召回和关键词召回结果，并按融合分排序。"""
        rrf_k = current_app.config.get('HYBRID_RRF_K', 60)
        merged = {}

        def add_candidate(doc, method):
            key = self._candidate_key(doc)
            incoming = dict(doc.metadata or {})
            if key not in merged:
                copied = LangChainDocument(
                    page_content=doc.page_content,
                    metadata=incoming
                )
                copied.metadata['retrieval_sources'] = [method]
                merged[key] = copied
                return

            existing = merged[key]
            metadata = existing.metadata
            sources = set(metadata.get('retrieval_sources', []))
            sources.add(method)
            metadata['retrieval_sources'] = sorted(sources)

            for field in (
                'vector_score',
                'keyword_score',
                'filename_score',
                'keyword_search_score',
                'bm25_score',
                'bm25_score_raw',
                'relevance_score',
            ):
                old_value = metadata.get(field)
                new_value = incoming.get(field)
                if isinstance(new_value, (int, float)):
                    if not isinstance(old_value, (int, float)) or new_value > old_value:
                        metadata[field] = new_value

            for field in ('vector_rank', 'keyword_rank'):
                old_rank = metadata.get(field)
                new_rank = incoming.get(field)
                if isinstance(new_rank, int):
                    if not isinstance(old_rank, int) or new_rank < old_rank:
                        metadata[field] = new_rank

            for field in (
                'keyword_backend',
                'bm25_query_terms',
                'bm25_matched_terms',
                'bm25_matched_keywords',
            ):
                if incoming.get(field) is not None:
                    metadata[field] = incoming[field]

            metadata['retrieval_method'] = '+'.join(metadata['retrieval_sources'])

        for doc in vector_docs:
            add_candidate(doc, 'vector')
        for doc in keyword_docs:
            add_candidate(doc, 'bm25')

        docs = list(merged.values())
        for doc in docs:
            metadata = doc.metadata
            hybrid_score = 0
            vector_rank = metadata.get('vector_rank')
            keyword_rank = metadata.get('keyword_rank')
            if isinstance(vector_rank, int):
                hybrid_score += 1 / (rrf_k + vector_rank)
            if isinstance(keyword_rank, int):
                hybrid_score += 1 / (rrf_k + keyword_rank)
            metadata['hybrid_score'] = round(hybrid_score, 6)
            metadata['hybrid_enabled'] = True

        docs.sort(key=lambda item: item.metadata.get('hybrid_score', 0), reverse=True)
        return docs

    def search(self, kb_id, query, top_k=None, use_rerank=None, candidate_k=None):
        """
        在指定知识库中检索相关文本块，并尽量保留相似度信息
        :param kb_id: 知识库ID
        :param query: 检索问题
        :param top_k: 返回数量
        :param use_rerank: 是否启用召回后重排
        :param candidate_k: 初始向量召回候选数量
        :return: 文档列表
        """
        vectorstore = self._get_vectorstore(kb_id)
        k = top_k or current_app.config['RETRIEVER_TOP_K']
        rerank_enabled = current_app.config.get('RERANK_ENABLED', True) if use_rerank is None else use_rerank
        initial_k = max(k, candidate_k or current_app.config.get('RERANK_CANDIDATE_K', k))
        hybrid_enabled = current_app.config.get('HYBRID_SEARCH_ENABLED', True)

        vector_docs = self._vector_search(vectorstore, query, initial_k)
        if hybrid_enabled:
            keyword_k = current_app.config.get('HYBRID_KEYWORD_CANDIDATE_K', initial_k)
            keyword_docs = self._keyword_search(vectorstore, query, keyword_k)
            docs = self._merge_hybrid_results(vector_docs, keyword_docs)
        else:
            docs = vector_docs
            for doc in docs:
                doc.metadata['hybrid_enabled'] = False

        if not rerank_enabled:
            for doc in docs[:k]:
                doc.metadata = dict(doc.metadata or {})
                doc.metadata['rerank_enabled'] = False
            return docs[:k]

        reranker = RerankService()
        return reranker.rerank(query, docs, k)

    def get_retriever(self, kb_id):
        """
        获取指定知识库的检索器
        :param kb_id: 知识库ID
        :return: Chroma检索器
        """
        vectorstore = self._get_vectorstore(kb_id)
        return vectorstore.as_retriever(
            search_kwargs={'k': current_app.config['RETRIEVER_TOP_K']}
        )
